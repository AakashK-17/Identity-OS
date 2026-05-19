import argparse
import json
import os
import re
import shutil
import subprocess
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from openai import OpenAI


# ---------------- CONFIG ---------------- #

BASE_RESUME = Path(r"D:\Resume's and coverletter\University of Utah Health Research\Aakash Kunarapu_Data Scientist_Template.docx")
OUTPUT_ROOT = Path(os.environ.get("OUTPUT_ROOT", r"D:\Resume's and coverletter"))
DEFAULT_JD_FILE = Path("job_description.txt")
DEFAULT_MODEL = "gpt-4o-mini"
MIN_GENERATED_WORDS = 520
FORBIDDEN_POSITIONING_WORDS = ["equivalent", "adjacent", "similar"]
FORMAT_REFERENCE = Path(__file__).parent / "format_reference_b.docx"
ANTI_AI_PATTERNS = [
    "stakeholder-ready",
    "recruiter-ready",
    "hiring-manager keyword proof",
    "interview-ready",
    "business-ready",
    "fast JD alignment",
    "semantic saturation",
    "keyword inventory",
]

PLACEHOLDERS = {
    "{{SUMMARY}}": "summary",
    "{{DESTINATION_CLEVELAND}}": "destination_cleveland_bullets",
    "{{GENPACT}}": "genpact_bullets",
    "{{PROJECTS}}": "projects",
    "{{CORE_COMPETENCIES}}": "core_competencies",
}

SECTION_ALIASES = {
    "summary": ["PROFESSIONAL SUMMARY", "SUMMARY", "PROFILE", "CAREER SUMMARY", "ABOUT"],
    "experience": ["EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT HISTORY"],
    "projects": ["PROJECTS", "PROJECT EXPERIENCE", "ACADEMIC PROJECTS", "SELECTED PROJECTS"],
    "competencies": ["CORE COMPETENCIES", "SKILLS", "TECHNICAL SKILLS", "CORE SKILLS", "COMPETENCIES"],
    "education": ["EDUCATION", "ACADEMIC BACKGROUND"],
    "certifications": ["CERTIFICATIONS", "CERTIFICATION", "LICENSES"],
}


# ---------------- TEXT HELPERS ---------------- #

def safe_filename(value: str, fallback: str, max_len: int = 80) -> str:
    value = (value or fallback).strip()
    value = re.sub(r"[<>:\"/\\|?*\n\r\t]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:max_len].strip(" .") or fallback


TITLE_KEYWORDS = [
    "accountant", "administrator", "advisor", "analyst", "architect", "associate", "consultant",
    "coordinator", "developer", "director", "engineer", "lead", "manager", "officer", "planner",
    "product manager", "program manager", "researcher", "scientist", "specialist", "strategist",
    "technician", "writer",
]
COMPANY_SUFFIX_PATTERN = (
    r"LLC|L\.L\.C\.|Inc\.?|Incorporated|Corporation|Corp\.?|Company|Co\.|Technologies|Technology|"
    r"Consulting|Agriscience|Health|Healthcare|University|College|Labs|Group|Bank|Capital|Global|"
    r"Systems|Solutions|Partners|Ventures|Media|Studios|Hospital|Clinic"
)
BAD_METADATA_PATTERNS = [
    r"school alumni from",
    r"people also viewed",
    r"promoted by",
    r"application (?:viewed|submitted|status)",
    r"take the next step",
    r"practice an interview",
    r"about the job",
    r"about the company",
    r"knowledge[, ]+skills",
    r"responsibilit(?:y|ies)",
    r"qualification",
    r"benefit",
    r"compensation",
    r"pay range",
    r"salary",
    r"equal opportunity",
    r"privacy",
    r"physical requirements",
    r"schedule",
    r"shift",
    r"address",
    r"city:",
    r"state:",
    r"postal code",
    r"remote|hybrid|onsite",
    r"full[- ]time|part[- ]time",
    r"over \d+ applicants",
    r"^[a-z][a-z .'-]+,\s*[a-z]{2}(?:\b|$)",
]
LINKEDIN_SEPARATOR_PATTERN = r"\s+(?:\||\?|\ufffd|\u00b7|\u2022)\s+"


def clean_metadata_value(value: str, max_len: int = 90) -> str:
    value = re.sub(r"\s+", " ", value or "").strip(" \t\r\n:,-|")
    value = re.sub(r"^(?:for|at|with)\s+", "", value, flags=re.IGNORECASE).strip()
    value = re.sub(r"\s+(?:required|preferred|not mandatory)$", "", value, flags=re.IGNORECASE).strip()
    return value[:max_len].strip(" .,-")


def clean_role_value(value: str) -> str:
    value = clean_metadata_value(value)
    value = re.sub(r"^(?:an?\s+|the\s+)?(?:highly skilled\s+|experienced\s+)?", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\s+(?:to|who|for|within|on)\s+.+$", "", value, flags=re.IGNORECASE).strip()
    return clean_metadata_value(value)


def metadata_lines(jd_text: str) -> list[str]:
    lines = []
    for raw in (jd_text or "").splitlines():
        line = clean_metadata_value(raw, max_len=180)
        if line:
            lines.append(line)
    return lines


def is_bad_metadata_line(line: str) -> bool:
    lower = (line or "").lower()
    if not lower or lower in {"role", "company", "job", "title", "description", "overview"}:
        return True
    return any(re.search(pattern, lower) for pattern in BAD_METADATA_PATTERNS)


def has_title_keyword(value: str) -> bool:
    lower = (value or "").lower()
    return any(re.search(rf"\b{re.escape(keyword)}s?\b", lower) for keyword in TITLE_KEYWORDS)


def is_probable_role(value: str) -> bool:
    lower = (value or "").lower()
    words = value.split()
    return (
        1 < len(words) <= 12
        and not re.search(LINKEDIN_SEPARATOR_PATTERN, value)
        and has_title_keyword(lower)
        and not is_bad_metadata_line(value)
    )


def is_probable_company(value: str) -> bool:
    cleaned = clean_metadata_value(value)
    lower = cleaned.lower()
    words = cleaned.split()
    if not cleaned or is_bad_metadata_line(cleaned) or len(words) > 10:
        return False
    if re.search(rf"\b(?:{COMPANY_SUFFIX_PATTERN})\b", cleaned, re.IGNORECASE):
        return True
    if re.search(r"\b[A-Z][A-Za-z]+(?:\.[A-Za-z]+)?\b", cleaned) and len(words) <= 5:
        return not any(keyword in lower for keyword in TITLE_KEYWORDS)
    return False


def labeled_metadata_value(lines: list[str], labels: list[str]) -> tuple[str, str] | None:
    label_pattern = "|".join(re.escape(label) for label in labels)
    for line in lines[:80]:
        match = re.match(rf"^\s*(?:{label_pattern})\s*[:\-]\s*(.+)$", line, re.IGNORECASE)
        if match:
            value = clean_metadata_value(match.group(1))
            if value and not is_bad_metadata_line(value):
                return value, f"label:{match.group(0).split(':', 1)[0].strip()}"
    return None


def linked_in_metadata(lines: list[str]) -> tuple[str, str] | None:
    for line in lines[:30]:
        if is_bad_metadata_line(line):
            continue
        parts = [clean_metadata_value(part) for part in re.split(LINKEDIN_SEPARATOR_PATTERN, line) if clean_metadata_value(part)]
        if len(parts) >= 2 and is_probable_role(parts[0]) and is_probable_company(parts[1]):
            return parts[0], parts[1]
    return None


def extract_role_from_text(jd_text: str, lines: list[str]) -> tuple[str, str] | None:
    labeled = labeled_metadata_value(lines, ["job title", "title", "role", "position", "job"])
    if labeled and is_probable_role(labeled[0]):
        return clean_role_value(labeled[0]), labeled[1]

    linkedin = linked_in_metadata(lines)
    if linkedin:
        return linkedin[0], "linkedin-style"

    top_lines = [line for line in lines[:25] if not is_bad_metadata_line(line)]
    for line in top_lines:
        if is_probable_role(line):
            return clean_role_value(re.sub(r"^(?:job title|title|role|position)\s*[:\-]\s*", "", line, flags=re.IGNORECASE)), "top-line"

    phrase_patterns = [
        r"\b(?:seeking|hiring|looking for|recruiting)\s+(?:an?\s+|the\s+)?(?:highly skilled\s+|experienced\s+)?([A-Z][A-Za-z0-9/&,+#.\- ]{2,90}?(?:%s)[A-Za-z0-9/&,+#.\- ]{0,60})(?:\s+to\b|\s+who\b|\.|,|\n)" % "|".join(TITLE_KEYWORDS),
        r"\bjoin (?:our|the)\s+.+?\s+as\s+(?:an?\s+)?([A-Z][A-Za-z0-9/&,+#.\- ]{2,90}?(?:%s)[A-Za-z0-9/&,+#.\- ]{0,60})(?:\.|,|\n)" % "|".join(TITLE_KEYWORDS),
    ]
    for pattern in phrase_patterns:
        match = re.search(pattern, jd_text, re.IGNORECASE)
        if match:
            role = clean_role_value(match.group(1))
            if is_probable_role(role):
                return role, "role-phrase"

    return None


def extract_company_from_text(jd_text: str, lines: list[str]) -> tuple[str, str] | None:
    labeled = labeled_metadata_value(lines, ["company", "organization", "employer", "client"])
    if labeled and is_probable_company(labeled[0]):
        return labeled

    logo_match = re.search(r"company logo for,?\s*([^\n\r.]+)", jd_text, re.IGNORECASE)
    if logo_match:
        company = clean_metadata_value(logo_match.group(1))
        if is_probable_company(company):
            return company, "company-logo"

    linkedin = linked_in_metadata(lines)
    if linkedin:
        return linkedin[1], "linkedin-style"

    company_patterns = [
        r"\b(?:Joining|At)\s+([A-Z][A-Za-z&.'+\- ]{2,70})\s+is\b",
        r"\b([A-Z][A-Za-z&.'+\- ]{2,70}?(?:%s))\s+is\s+(?:a|an|the)\b" % COMPANY_SUFFIX_PATTERN,
        r"\b([A-Z][A-Za-z&.'+\- ]{2,70}?(?:%s))\s+(?:is seeking|is hiring|invites applications|seeks)\b" % COMPANY_SUFFIX_PATTERN,
        r"\bThrough its division\s+([A-Z][A-Za-z&.'+\- ]{2,70})\b",
    ]
    for pattern in company_patterns:
        match = re.search(pattern, jd_text, re.IGNORECASE)
        if match:
            company = clean_metadata_value(match.group(1))
            if is_probable_company(company):
                return company, "company-phrase"

    for line in lines[:20]:
        if is_probable_company(line):
            return clean_metadata_value(line), "top-line"

    return None


def extract_job_metadata(jd_text: str) -> dict:
    lines = metadata_lines(jd_text)
    role_result = extract_role_from_text(jd_text, lines)
    company_result = extract_company_from_text(jd_text, lines)

    role_display = role_result[0] if role_result else "Target Role"
    company_display = company_result[0] if company_result else "Unknown Company"
    confidence = 0
    if company_result:
        confidence += 50
    if role_result:
        confidence += 50

    return {
        "company_display": company_display,
        "role_display": role_display,
        "company_filename": safe_filename(company_display, "Unknown Company"),
        "role_filename": safe_filename(role_display, "Target Role"),
        "metadata_confidence": confidence,
        "metadata_source": {
            "company": company_result[1] if company_result else "fallback",
            "role": role_result[1] if role_result else "fallback",
        },
    }


def extract_company_role(jd_text: str) -> tuple[str, str]:
    metadata = extract_job_metadata(jd_text)
    return metadata["company_display"], metadata["role_display"]


def extract_json(text: str) -> dict:
    """Handle plain JSON or JSON wrapped in a Markdown code fence."""
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)```", text, re.DOTALL | re.IGNORECASE)
    if fenced:
        text = fenced.group(1).strip()
    return json.loads(text)


def flatten_generated_text(data: dict) -> str:
    parts = [str(data.get("summary", ""))]
    parts.extend(str(item) for item in data.get("destination_cleveland_bullets", []))
    parts.extend(str(item) for item in data.get("genpact_bullets", []))

    for project in data.get("projects", []):
        if isinstance(project, dict):
            parts.append(str(project.get("title", "")))
            parts.extend(str(item) for item in project.get("bullets", []))
        else:
            parts.append(str(project))

    parts.extend(str(item) for item in data.get("core_competencies", []))
    return "\n".join(parts)


def generated_word_count(data: dict) -> int:
    text = re.sub(r"\*\*", "", flatten_generated_text(data))
    return len(re.findall(r"\b[\w+#./-]+\b", text))


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w+#./-]+\b", re.sub(r"\*\*", "", str(text))))


def expand_sentence(text: str, target_min: int, filler: str) -> str:
    text = str(text).strip().rstrip(".")
    if word_count(text) >= target_min:
        return text + "."

    filler_words = filler.strip().rstrip(".")
    separator = " using " if not re.search(r"\busing\b", text, re.IGNORECASE) else " for "
    expanded = f"{text}{separator}{filler_words}"
    return expanded.strip().rstrip(".") + "."


def trim_to_max_words(text: str, max_words: int) -> str:
    words = str(text).strip().split()
    if len(words) <= max_words:
        return str(text).strip()
    return " ".join(words[:max_words]).rstrip(".,;") + "."


def repair_generated_resume(data: dict, jd_text: str) -> dict:
    """Lightly repair shape issues without fabricating tools, domains, or metrics."""
    data.setdefault("summary", "")
    data.setdefault("destination_cleveland_bullets", [])
    data.setdefault("genpact_bullets", [])
    data.setdefault("projects", [])
    data.setdefault("core_competencies", [])

    if word_count(data.get("summary", "")) > 90:
        data["summary"] = trim_to_max_words(data.get("summary", ""), 82)

    for key in ["destination_cleveland_bullets", "genpact_bullets"]:
        bullets = data.get(key, [])
        data[key] = [trim_to_max_words(str(bullet), 32) for bullet in bullets[:5]]

    for project_index, project in enumerate(data.get("projects", [])):
        if not isinstance(project, dict):
            continue
        project["bullets"] = [trim_to_max_words(str(bullet), 36) for bullet in project.get("bullets", [])[:3]]

    if len(data.get("core_competencies", [])) < 6:
        data.setdefault("core_competencies", []).append(
            "Professional Strengths: problem-solving, collaboration, communication, documentation, continuous learning"
        )

    for word in FORBIDDEN_POSITIONING_WORDS:
        pattern = re.compile(rf"\b{re.escape(word)}\b", re.IGNORECASE)
        for key in ["summary", "destination_cleveland_bullets", "genpact_bullets", "core_competencies"]:
            if isinstance(data.get(key), str):
                data[key] = pattern.sub("", data[key])
            elif isinstance(data.get(key), list):
                data[key] = [pattern.sub("", str(item)).replace("  ", " ").strip() for item in data[key]]

        for project in data.get("projects", []):
            if isinstance(project, dict):
                project["title"] = pattern.sub("", str(project.get("title", ""))).replace("  ", " ").strip()
                project["bullets"] = [
                    pattern.sub("", str(item)).replace("  ", " ").strip()
                    for item in project.get("bullets", [])
                ]

    for phrase in ANTI_AI_PATTERNS:
        pattern = re.compile(re.escape(phrase), re.IGNORECASE)
        for key in ["summary", "destination_cleveland_bullets", "genpact_bullets", "core_competencies"]:
            if isinstance(data.get(key), str):
                data[key] = pattern.sub("", data[key]).replace("  ", " ").strip()
            elif isinstance(data.get(key), list):
                data[key] = [pattern.sub("", str(item)).replace("  ", " ").strip() for item in data[key]]

        for project in data.get("projects", []):
            if isinstance(project, dict):
                project["title"] = pattern.sub("", str(project.get("title", ""))).replace("  ", " ").strip()
                project["bullets"] = [
                    pattern.sub("", str(item)).replace("  ", " ").strip()
                    for item in project.get("bullets", [])
                ]

    return data


def validate_generated_resume(data: dict) -> list[str]:
    issues = []
    word_count = generated_word_count(data)
    if word_count < MIN_GENERATED_WORDS:
        issues.append(
            f"Generated content is thin at {word_count} words. Add truthful, interview-defensible detail from the provided profile."
        )

    summary_words = len(re.findall(r"\b[\w+#./-]+\b", str(data.get("summary", ""))))
    if summary_words < 45 or summary_words > 90:
        issues.append(f"Summary should be 45-90 words and identity-driven; current summary is {summary_words} words.")

    summary_lower = str(data.get("summary", "")).lower()
    if summary_lower.startswith("data scientist with"):
        issues.append("Summary uses the banned generic opening 'Data Scientist with'. Make it identity-driven and role-native.")

    for key in ["destination_cleveland_bullets", "genpact_bullets"]:
        bullets = data.get(key, [])
        if len(bullets) != 5:
            issues.append(f"{key} must contain exactly 5 bullets.")
        for index, bullet in enumerate(bullets, start=1):
            count = len(re.findall(r"\b[\w+#./-]+\b", str(bullet)))
            if count < 12 or count > 34:
                issues.append(f"{key} bullet {index} should be concise but natural; current count is {count}.")

    projects = data.get("projects", [])
    if len(projects) != 3:
        issues.append("projects must contain exactly 3 project objects.")
    for index, project in enumerate(projects, start=1):
        if not isinstance(project, dict):
            issues.append(f"Project {index} must be an object with title and bullets.")
            continue
        bullets = project.get("bullets", [])
        if len(bullets) < 2 or len(bullets) > 3:
            issues.append(f"Project {index} must contain 2-3 bullets.")
        for bullet_index, bullet in enumerate(bullets, start=1):
            count = len(re.findall(r"\b[\w+#./-]+\b", str(bullet)))
            if count < 14 or count > 40:
                issues.append(f"Project {index} detail {bullet_index} should be readable and defensible; current count is {count}.")

    competencies = data.get("core_competencies", [])
    if len(competencies) < 5 or len(competencies) > 6:
        issues.append("core_competencies must contain 5-6 categories.")

    lowered = flatten_generated_text(data).lower()
    for word in FORBIDDEN_POSITIONING_WORDS:
        if re.search(rf"\b{re.escape(word)}\b", lowered):
            issues.append(f"Do not use the forbidden positioning word: {word}.")

    for pattern in ANTI_AI_PATTERNS:
        if pattern in lowered:
            issues.append(f"Remove synthetic phrase: {pattern}.")

    repeated_starts = {}
    for bullet in data.get("destination_cleveland_bullets", []) + data.get("genpact_bullets", []):
        first = str(bullet).strip().split(" ", 1)[0].lower()
        if first:
            repeated_starts[first] = repeated_starts.get(first, 0) + 1
    for first, count in repeated_starts.items():
        if count > 2:
            issues.append(f"Too many bullets start with '{first}'. Vary cadence.")

    return issues


def validate_authenticity(data: dict, base_resume_text: str) -> list[str]:
    return []


def strip_unsupported_terms(data: dict, base_resume_text: str) -> dict:
    return data


def make_template_from_resume(source_path: Path, template_path: Path) -> Path:
    """Create a strict placeholder template from a resume with the expected section headings."""
    doc = Document(str(source_path))

    summary_heading = paragraph_index(doc, "PROFESSIONAL SUMMARY", exact=True)
    experience_heading = paragraph_index(doc, "EXPERIENCE", exact=True)
    replace_range_after_anchor(doc, summary_heading, experience_heading, ["{{SUMMARY}}"])

    dc_heading = paragraph_index(doc, "Destination Cleveland")
    genpact_heading = paragraph_index(doc, "Genpact", start=max(dc_heading, 0))
    replace_range_after_anchor(doc, dc_heading + 1, genpact_heading, ["{{DESTINATION_CLEVELAND}}"])

    genpact_heading = paragraph_index(doc, "Genpact")
    projects_heading = paragraph_index(doc, "PROJECTS", start=max(genpact_heading, 0), exact=True)
    replace_range_after_anchor(doc, genpact_heading + 1, projects_heading, ["{{GENPACT}}"])

    projects_heading = paragraph_index(doc, "PROJECTS", exact=True)
    core_heading = paragraph_index(doc, "CORE COMPETENCIES", start=max(projects_heading, 0), exact=True)
    replace_range_after_anchor(doc, projects_heading, core_heading, ["{{PROJECTS}}"])

    core_heading = paragraph_index(doc, "CORE COMPETENCIES", exact=True)
    education_heading = paragraph_index(doc, "EDUCATION", start=max(core_heading, 0), exact=True)
    replace_range_after_anchor(doc, core_heading, education_heading, ["{{CORE_COMPETENCIES}}"])

    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def set_bottom_border(paragraph, color="000000", size="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def create_template_from_profile(profile: dict, template_path: Path) -> Path:
    doc = Document(str(FORMAT_REFERENCE)) if FORMAT_REFERENCE.exists() else Document()

    details = profile.get("details", {})
    name = details.get("name") or "Candidate Name"
    has_contact = any(
        (details.get(key) or "").strip()
        for key in ["location", "phone", "email", "linkedin"]
    )

    if not FORMAT_REFERENCE.exists() or len(doc.paragraphs) < 10:
        # Emergency fallback only. Normal path keeps the user's B-format skeleton.
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(name)
        name_run.bold = True
        name_paragraph.alignment = 1
        if contact:
            contact_paragraph = doc.add_paragraph(contact)
            contact_paragraph.alignment = 1
        doc.add_paragraph("PROFESSIONAL SUMMARY").style = "Heading 1"
        doc.add_paragraph("{{SUMMARY}}")
        doc.add_paragraph("EXPERIENCE").style = "Heading 1"
        doc.add_paragraph("{{DESTINATION_CLEVELAND}}")
        doc.add_paragraph("{{GENPACT}}")
        doc.add_paragraph("PROJECTS").style = "Heading 1"
        doc.add_paragraph("{{PROJECTS}}")
        doc.add_paragraph("CORE COMPETENCIES").style = "Heading 1"
        doc.add_paragraph("{{CORE_COMPETENCIES}}")
    else:
        # Preserve the exact B-format paragraph structure, spacing, styles, and margins.
        while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
            delete_paragraph(doc.paragraphs[-1])

        set_paragraph_text_preserve_first_run(doc.paragraphs[0], name)
        if contact and len(doc.paragraphs) > 1:
            set_paragraph_text_preserve_first_run(doc.paragraphs[1], contact)

        summary_idx = paragraph_index(doc, "PROFESSIONAL SUMMARY", exact=True)
        exp_idx = paragraph_index(doc, "EXPERIENCE", exact=True)
        projects_idx = paragraph_index(doc, "PROJECTS", exact=True)
        competencies_idx = paragraph_index(doc, "CORE COMPETENCIES", exact=True)
        education_idx = paragraph_index(doc, "EDUCATION", exact=True)
        certifications_idx = paragraph_index(doc, "CERTIFICATIONS", exact=True)

        if summary_idx >= 0 and exp_idx > summary_idx:
            replace_range_after_anchor(doc, summary_idx, exp_idx, ["{{SUMMARY}}"])

        experiences = profile.get("experiences", [])[:2]
        while len(experiences) < 2:
            experiences.append({})

        experience_lines = []
        for index, experience in enumerate(experiences):
            company = experience.get("company") or f"Company {index + 1}"
            title = experience.get("title") or "Role"
            duration = experience.get("duration") or ""
            location = experience.get("location") or ""
            header = f"{company} – {title}"
            if duration:
                header += f"\t{duration}"
            experience_lines.append(header)
            if location:
                experience_lines.append(location)
            experience_lines.append("{{DESTINATION_CLEVELAND}}" if index == 0 else "{{GENPACT}}")

        projects_idx = paragraph_index(doc, "PROJECTS", exact=True)
        exp_idx = paragraph_index(doc, "EXPERIENCE", exact=True)
        if exp_idx >= 0 and projects_idx > exp_idx:
            replace_range_after_anchor(doc, exp_idx, projects_idx, experience_lines)

        competencies_idx = paragraph_index(doc, "CORE COMPETENCIES", exact=True)
        projects_idx = paragraph_index(doc, "PROJECTS", exact=True)
        if projects_idx >= 0 and competencies_idx > projects_idx:
            replace_range_after_anchor(doc, projects_idx, competencies_idx, ["{{PROJECTS}}"])

        education_idx = paragraph_index(doc, "EDUCATION", exact=True)
        competencies_idx = paragraph_index(doc, "CORE COMPETENCIES", exact=True)
        if competencies_idx >= 0 and education_idx > competencies_idx:
            replace_range_after_anchor(doc, competencies_idx, education_idx, ["{{CORE_COMPETENCIES}}"])

        education = profile.get("education", [])
        certifications_idx = paragraph_index(doc, "CERTIFICATIONS", exact=True)
        education_idx = paragraph_index(doc, "EDUCATION", exact=True)
        if education_idx >= 0:
            education_lines = []
            for item in education:
                school = item.get("school", "")
                degree = item.get("degree", "")
                year = item.get("year", "")
                line = " — ".join(part for part in [school, degree] if part)
                if year:
                    line += f"\t{year}"
                if line.strip():
                    education_lines.append(line)
            education_end = certifications_idx if certifications_idx > education_idx else len(doc.paragraphs)
            replace_range_after_anchor(doc, education_idx, education_end, education_lines)

        certifications = profile.get("certifications", [])
        certifications_idx = paragraph_index(doc, "CERTIFICATIONS", exact=True)
        if certifications_idx >= 0:
            certification_lines = []
            for item in certifications:
                text = item.get("name", "") if isinstance(item, dict) else str(item)
                if text.strip():
                    certification_lines.append(text.strip())
            replace_range_after_anchor(doc, certifications_idx, len(doc.paragraphs), certification_lines)

        # Bold experience headers created from profile data, matching B's visual language.
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if " – " in text and "\t" in text and paragraph.runs:
                for run in paragraph.runs:
                    run.bold = True

    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def create_template_from_profile(profile: dict, template_path: Path) -> Path:
    """Build the screenshot-style resume structure from stored profile fields."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.06

    details = profile.get("details", {})
    name = details.get("name") or "Candidate Name"
    contact = " • ".join(
        value for value in [
            details.get("location", ""),
            details.get("phone", ""),
            details.get("email", ""),
            details.get("linkedin", ""),
        ]
        if value
    )
    has_contact = any(
        (details.get(key) or "").strip()
        for key in ["location", "phone", "email", "linkedin"]
    )

    def set_spacing(paragraph, before=0, after=2.2, line_spacing=1.06):
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = line_spacing
        return paragraph

    def add_run(paragraph, text, bold=False, italic=False, size=9.6):
        run = paragraph.add_run(str(text))
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return run

    def add_plain(text="", bold=False, italic=False, size=9.6, align=None, after=2.2):
        paragraph = doc.add_paragraph()
        set_spacing(paragraph, after=after)
        if align is not None:
            paragraph.alignment = align
        add_run(paragraph, text, bold=bold, italic=italic, size=size)
        return paragraph

    def add_section_heading(text):
        paragraph = add_plain(text, bold=True, size=11.6, after=2)
        set_bottom_border(paragraph)
        return paragraph

    def add_tabbed_line(left, right="", bold=False, italic=False):
        paragraph = doc.add_paragraph()
        set_spacing(paragraph, after=1.4)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.28), WD_TAB_ALIGNMENT.RIGHT)
        add_run(paragraph, left, bold=bold, italic=italic)
        if right:
            add_run(paragraph, "\t" + str(right), bold=bold)
        return paragraph

    add_plain(name, bold=True, size=9.8, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    if has_contact:
        contact_paragraph = add_plain("", size=9.6, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
        set_contact_line(contact_paragraph, details)

    add_section_heading("PROFESSIONAL SUMMARY")
    add_plain("{{SUMMARY}}", after=6)

    add_section_heading("EXPERIENCE")
    experiences = profile.get("experiences", [])[:2]
    while len(experiences) < 2:
        experiences.append({})

    for index, experience in enumerate(experiences):
        company = experience.get("company") or f"Company {index + 1}"
        title = experience.get("title") or "Role"
        duration = experience.get("duration") or ""
        location = experience.get("location") or ""
        add_tabbed_line(f"{company} – {title}", duration, bold=True)
        if location:
            add_plain(location, italic=True, after=1)
        placeholder = "{{DESTINATION_CLEVELAND}}" if index == 0 else "{{GENPACT}}"
        add_plain(placeholder, after=7 if index == 0 else 11)

    add_section_heading("PROJECTS")
    add_plain("{{PROJECTS}}", after=10)

    add_section_heading("CORE COMPETENCIES")
    add_plain("{{CORE_COMPETENCIES}}", after=8)

    education = profile.get("education", [])
    if education:
        add_section_heading("EDUCATION")
        for item in education:
            school = item.get("school", "")
            degree = item.get("degree", "")
            year = item.get("year", "")
            if school:
                add_tabbed_line(school, year)
            if degree:
                add_plain(degree, italic=True, after=1.2)

    certifications = profile.get("certifications", [])
    if certifications:
        add_section_heading("CERTIFICATIONS")
        for item in certifications:
            text = item.get("name", "") if isinstance(item, dict) else str(item)
            if text.strip():
                add_plain(text.strip(), after=1)

    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))
    return template_path


def document_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def update_contact_details(doc: Document, details: dict) -> None:
    name = (details.get("name") or "").strip()
    location = (details.get("location") or "").strip()
    phone = (details.get("phone") or "").strip()
    email = (details.get("email") or "").strip()
    linkedin = (details.get("linkedin") or "").strip()

    if name and doc.paragraphs:
        set_paragraph_text_preserve_first_run(doc.paragraphs[0], name)

    contact_parts = [part for part in [location, phone, email, linkedin] if part]
    if contact_parts and len(doc.paragraphs) > 1:
        set_paragraph_text_preserve_first_run(doc.paragraphs[1], " • ".join(contact_parts))


def set_paragraph_text_preserve_first_run(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    run_properties.append(underline)

    bold = OxmlElement("w:b")
    run_properties.append(bold)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def normalize_linkedin_url(value: str) -> str:
    value = (value or "").strip()
    if not value or value.lower() == "linkedin":
        return ""
    if value.startswith("http://") or value.startswith("https://"):
        return value
    if "linkedin.com" in value.lower():
        return "https://" + value.lstrip("/")
    return ""


def set_contact_line(paragraph, details: dict) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    parts = [
        (details.get("location") or "").strip(),
        (details.get("phone") or "").strip(),
        (details.get("email") or "").strip(),
    ]
    parts = [part for part in parts if part]
    linkedin_url = normalize_linkedin_url(details.get("linkedin", ""))

    for index, part in enumerate(parts):
        if index:
            paragraph.add_run(" • ")
        paragraph.add_run(part)

    if linkedin_url or (details.get("linkedin") or "").strip():
        if parts:
            paragraph.add_run(" • ")
        if linkedin_url:
            add_hyperlink(paragraph, "LinkedIn", linkedin_url)
        else:
            paragraph.add_run("LinkedIn")

    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(9.6)


def update_contact_details(doc: Document, details: dict) -> None:
    name = (details.get("name") or "").strip()
    clean_details = {
        "location": (details.get("location") or "").strip(),
        "phone": (details.get("phone") or "").strip(),
        "email": (details.get("email") or "").strip(),
        "linkedin": (details.get("linkedin") or "").strip(),
    }

    if name and doc.paragraphs:
        set_paragraph_text_preserve_first_run(doc.paragraphs[0], name)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if any(clean_details.values()) and len(doc.paragraphs) > 1:
        set_contact_line(doc.paragraphs[1], clean_details)


def format_value(value) -> str:
    if isinstance(value, list):
        lines = []
        for item in value:
            if isinstance(item, dict):
                title = str(item.get("title", "")).strip()
                bullets = item.get("bullets", [])
                body = " ".join(str(bullet).strip().rstrip(".") + "." for bullet in bullets if str(bullet).strip())
                if title and body:
                    lines.append(f"{title}: {body}")
                elif title:
                    lines.append(title)
            else:
                lines.append(str(item).strip())
        return "\n".join(line for line in lines if line)
    return str(value).strip()


def add_text_with_bold(paragraph, text: str) -> None:
    paragraph.clear()
    parts = re.split(r"(\*\*.*?\*\*)", str(text))
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        run.font.size = Pt(9.6)
        run.bold = part.startswith("**") and part.endswith("**")


def apply_bullet(paragraph) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.hanging_indent = None


def replace_inline_placeholder(paragraph, placeholder: str, new_text: str) -> bool:
    if placeholder not in paragraph.text:
        return False

    add_text_with_bold(paragraph, paragraph.text.replace(placeholder, new_text))
    return True


def insert_formatted_paragraph_after(paragraph, text: str, style=None) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    source_p_pr = paragraph._p.pPr
    if source_p_pr is not None:
        new_paragraph._p.insert(0, deepcopy(source_p_pr))
    if style:
        try:
            new_paragraph.style = style
        except KeyError:
            pass
        apply_bullet(new_paragraph)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text_with_bold(new_paragraph, text)
    return new_paragraph


def replace_block_placeholder(doc: Document, placeholder: str, value, style=None) -> bool:
    if placeholder == "{{PROJECTS}}" and isinstance(value, list):
        replacement_lines = []
        for project in value:
            if isinstance(project, dict):
                title = str(project.get("title", "")).strip()
                bullets = [str(item).strip() for item in project.get("bullets", []) if str(item).strip()]
                body = " ".join(bullet.rstrip(".") + "." for bullet in bullets)
                if title and body:
                    replacement_lines.append({"text": f"**{title}:** {body}", "bullet": False})
                elif title:
                    replacement_lines.append({"text": f"**{title}**", "bullet": False})
                else:
                    replacement_lines.append({"text": body, "bullet": False})
            else:
                replacement_lines.append({"text": str(project).strip(), "bullet": False})
    else:
        raw_lines = value if isinstance(value, list) else [value]
        replacement_lines = []
        for item in raw_lines:
            text = str(item).strip()
            if placeholder == "{{CORE_COMPETENCIES}}" and ":" in text and "**" not in text.split(":", 1)[0]:
                key, rest = text.split(":", 1)
                text = f"**{key.strip()}:**{rest}"
            if style is not None and text and not text.startswith("•"):
                text = "• " + text
            replacement_lines.append({"text": text, "bullet": False})

    replacement_lines = [item for item in replacement_lines if item["text"]]

    for paragraph in list(doc.paragraphs):
        if placeholder not in paragraph.text:
            continue

        if paragraph.text.strip() == placeholder:
            for line in reversed(replacement_lines):
                line_style = style if line["bullet"] else None
                insert_formatted_paragraph_after(paragraph, line["text"], style=line_style)
            delete_paragraph(paragraph)
            return True

        return replace_inline_placeholder(paragraph, placeholder, format_value(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if placeholder not in paragraph.text:
                        continue

                    if paragraph.text.strip() == placeholder:
                        for line in reversed(replacement_lines):
                            line_style = style if line["bullet"] else None
                            insert_formatted_paragraph_after(paragraph, line["text"], style=line_style)
                        delete_paragraph(paragraph)
                        return True

                    return replace_inline_placeholder(paragraph, placeholder, format_value(value))

    return False


def replace_placeholders(doc: Document, data: dict) -> list[str]:
    missing = []
    for placeholder, key in PLACEHOLDERS.items():
        value = data.get(key, "")
        if key in {"destination_cleveland_bullets", "genpact_bullets", "core_competencies"} and isinstance(value, list):
            value = [item if str(item).strip().startswith("•") else f"• {item}" for item in value]
        found = replace_block_placeholder(doc, placeholder, value, style=None)
        if not found:
            missing.append(placeholder)
    return missing


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraphs_after(paragraph, lines: list[str], style=None):
    current = paragraph
    for line in lines:
        new_element = OxmlElement("w:p")
        current._p.addnext(new_element)
        current = Paragraph(new_element, current._parent)
        current.text = str(line).strip()
        if style:
            current.style = style
    return current


def paragraph_index(doc: Document, text: str, start: int = 0, exact: bool = False) -> int:
    needle = text.lower()
    for idx, paragraph in enumerate(doc.paragraphs[start:], start=start):
        haystack = paragraph.text.strip().lower()
        if exact and haystack == needle:
            return idx
        if not exact and needle in haystack:
            return idx
    return -1


def replace_range_after_anchor(doc: Document, anchor_idx: int, end_idx: int, lines: list[str]) -> None:
    if anchor_idx < 0 or end_idx <= anchor_idx:
        return

    anchor = doc.paragraphs[anchor_idx]
    old_paragraphs = list(doc.paragraphs[anchor_idx + 1:end_idx])

    for paragraph in old_paragraphs:
        delete_paragraph(paragraph)

    insert_paragraphs_after(anchor, lines)


def replace_existing_resume_sections(doc: Document, data: dict) -> None:
    """Fallback for the current base resume when placeholders have not been added."""
    summary_heading = paragraph_index(doc, "PROFESSIONAL SUMMARY", exact=True)
    experience_heading = paragraph_index(doc, "EXPERIENCE", exact=True)
    if summary_heading >= 0 and experience_heading > summary_heading:
        replace_range_after_anchor(
            doc,
            summary_heading,
            experience_heading,
            [format_value(data.get("summary", ""))],
        )

    dc_heading = paragraph_index(doc, "Destination Cleveland")
    genpact_heading = paragraph_index(doc, "Genpact", start=max(dc_heading, 0))
    projects_heading = paragraph_index(doc, "PROJECTS", start=max(genpact_heading, 0), exact=True)

    if dc_heading >= 0 and genpact_heading > dc_heading:
        # Keep the location line, replace only the experience bullets.
        replace_range_after_anchor(
            doc,
            dc_heading + 1,
            genpact_heading,
            data.get("destination_cleveland_bullets", []),
        )

    genpact_heading = paragraph_index(doc, "Genpact")
    projects_heading = paragraph_index(doc, "PROJECTS", start=max(genpact_heading, 0), exact=True)
    if genpact_heading >= 0 and projects_heading > genpact_heading:
        replace_range_after_anchor(
            doc,
            genpact_heading + 1,
            projects_heading,
            data.get("genpact_bullets", []),
        )

    projects_heading = paragraph_index(doc, "PROJECTS", exact=True)
    skills_heading = paragraph_index(doc, "CORE COMPETENCIES", start=max(projects_heading, 0), exact=True)
    if projects_heading >= 0 and skills_heading > projects_heading:
        project_lines = []
        for project in data.get("projects", []):
            if isinstance(project, dict):
                details = " ".join(str(item).strip() for item in project.get("bullets", []) if str(item).strip())
                project_lines.append(f"{project.get('title', '').strip()}: {details}".strip(": "))
            elif str(project).strip():
                project_lines.append(str(project).strip())
        replace_range_after_anchor(doc, projects_heading, skills_heading, project_lines)

    skills_heading = paragraph_index(doc, "CORE COMPETENCIES", exact=True)
    education_heading = paragraph_index(doc, "EDUCATION", start=max(skills_heading, 0), exact=True)
    if skills_heading >= 0 and education_heading > skills_heading:
        competencies = data.get("core_competencies", [])
        if isinstance(competencies, list):
            competency_lines = [str(item).strip() for item in competencies if str(item).strip()]
        else:
            competency_lines = [str(competencies).strip()]
        replace_range_after_anchor(doc, skills_heading, education_heading, competency_lines)


# ---------------- LLM ---------------- #

def build_prompt(base_resume_text: str, jd_text: str) -> str:
    return f"""
You are an elite resume strategist and hiring committee analyst.

Use the Universal Resume Shortlisting Framework:
1. Deconstruct the JD into:
   - Core Functional Requirements
   - Systems / Infrastructure Signals
   - Risk Signals, meaning what the employer is afraid of
   - Tooling Stack
   - Cultural / Domain Alignment
   - Seniority Level
2. Identify gaps and positioning opportunities in the current resume.
3. Choose one best positioning persona:
   Data Infrastructure Builder, Analytical Problem Solver, Research Operations Partner,
   Systems Optimizer, Insight Translator, Technical Architect, or Growth Analyst.
4. Rewrite:
   - Professional Summary
   - Experience bullets
   - Projects
   - Core Competencies
5. Use the bullet formula:
   Action + System/Method + Context + Measurable Impact + Risk Reduction.
6. Ensure:
   - Natural keyword mirroring for ATS
   - Human readability
   - No exaggeration
   - Clear measurable outcomes where available
   - Reproducibility and structure emphasis when relevant
7. The resume should reduce perceived hiring risk and sound like the candidate has already solved the employer's problem.

Built-in generation rule:
Follow the framework and make sure important JD keywords are present and proven in Experience, Projects, or Core Competencies.
If the current experience is not a literal match for the JD, find truthful scope inside the candidate's experience that can be reframed to meet the JD requirements and company needs.
Improve readability and do not write unnecessarily long lines.
Strict note: do not use words like "equivalent", "adjacent", or "similar".
If the JD says AWS, include AWS along with relevant services such as S3, EC2, Lambda, Glue, Redshift, SageMaker, or CloudWatch when they fit the candidate's background.
The candidate accepts a resume going over one page, but does not accept a resume that fails to fill one complete page.

Critical automation instruction:
Rewrite every resume section from Professional Summary through Core Competencies.
Do not partially edit the resume.
Maintain original formatting from the uploaded/template DOCX.
Use bullet points in Experience.
Bold important JD keywords naturally using **keyword** markdown.
Core Competencies must use paired category format:
• Category: keyword, keyword, keyword.
Remove duplicate project text.
Keep readability high.
Every word must earn its place. Keep keywords, proof points, and differentiators. Remove filler.
Do not underwrite. The generated sections must contain at least {MIN_GENERATED_WORDS} total words before formatting.
Final trust test before returning JSON:
Would a recruiter believe this candidate exists and want to interview them?
If not, rewrite the weak sections before returning JSON.

Strict output requirements:
- Return valid JSON only.
- No Markdown fences.
- No commentary outside JSON.
- Rewrite the whole resume body, not isolated fragments.
- summary must be one paragraph of 60-80 words and should visually read as 4-6 lines in a resume.
- summary must include: positioning identity, years of experience and degree, top 3-4 JD-matched skills, and one differentiator.
- destination_cleveland_bullets must contain exactly 5 bullets of 15-25 words each.
- genpact_bullets must contain exactly 5 bullets of 15-25 words each.
- Experience bullets must use: Action verb + what was done + tool/method + result.
- No stacking clauses with commas.
- Do not use "and" to connect two different actions in one bullet.
- If a bullet needs a semicolon, split the idea before returning JSON.
- projects must contain exactly 3 project objects.
- Each project object must have a title with tools in parentheses and 2-3 detail strings.
- The DOCX will render each project as one paragraph, not separate project bullets.
- Each project paragraph must quickly show tools, what was built, what it did, and one measurable outcome if available.
- core_competencies must contain 5-6 paired category bullets.
- Each core competency category must contain 4-8 keywords.
- Categories should map to JD signals: core functional skills, technical tools, domain methods, communication/stakeholders, data/reporting tools, and culture/professional attributes only if signaled.
- Keywords inside each category should mirror the JD phrasing where truthful.
- Do not include bullet symbols in JSON values; the DOCX writer adds bullets.
- Use **bold** around important ATS keywords, tools, domain phrases, and differentiators.
- Target 30-50 bolded terms across the entire generated resume.
- Include Technical Skills inside core_competencies when useful.
- Do not make generic claims like "increased accuracy" unless grounded in the source resume or clearly framed as project output.
- If a JD keyword is important, prove it in Experience or Projects, not only in Core Competencies.
- Never return a thin resume. Expand Projects and Core Competencies first, then Experience bullets within the allowed range.

JSON schema:
{{
  "summary": "string",
  "destination_cleveland_bullets": ["bullet", "bullet", "bullet", "bullet", "bullet"],
  "genpact_bullets": ["bullet", "bullet", "bullet", "bullet", "bullet"],
  "projects": [
    {{
      "title": "Project Name (Python, SQL, Power BI)",
      "bullets": ["bullet", "bullet"]
    }}
  ],
  "core_competencies": [
    "Machine Learning & AI: Python, scikit-learn, RAG, vector databases, evaluation metrics",
    "Data Analysis & Statistics: SQL, Pandas, NumPy, hypothesis testing, regression"
  ]
}}

BASE RESUME:
{base_resume_text}

JOB DESCRIPTION:
{jd_text}
""".strip()


def build_prompt(base_resume_text: str, jd_text: str) -> str:
    return f"""
You are an elite resume strategist, recruiter-skeptic, and hiring committee analyst.

Use The ATS-Dominance & Contextual Alignment Protocol v4.0-X:
High-ATS, recruiter-readable resume architecture.

Primary objective:
Maximize ATS alignment and recruiter relevance while preserving readable, coherent storytelling.
Target 95%+ ATS alignment through exact phrase matching, weighted repetition, and strong section placement.

Core principle:
The resume is a bridge, not a mirror. Bridge the JD and the base profile through structural plausibility, transferable workflows, and exact JD vocabulary.
Experience is the credibility anchor. Projects are the capability expansion zone. Core Competencies are parser support.

PHASE 1 - Keyword Extraction & Intent Mapping:
- Vocabulary harvest: identify hard skills, tools, cloud platforms, data stores, methods, domain terms, and preferred qualifications.
- Operational intent: determine the JD's technical north star, such as ML scaling, personalization, analytics automation, cloud workflows, data infrastructure, or stakeholder reporting.
- Proof logic: for every important JD signal, find the closest usable hook in experience or projects and fit the exact JD phrase into the generated resume.
- Mirror the employer's exact phrase when possible. If the JD says "customer segmentation", do not replace it with "audience grouping".

PHASE 2 - Contextual Alignment Engine:
- Use exact keyword injection across Summary, Experience, Projects, and Core Competencies.
- Important terms should appear 2-4 times naturally when the JD plan requests it; lower-priority terms usually need 1-2 mentions.
- Prefer high-value placement in Summary, recent Experience, Projects, and Core Competencies.
- Build semantic-density clusters such as SQL + segmentation + forecasting + dashboards + experimentation so the resume reads deeply aligned, not merely keyword sprinkled.
- If experience does not naturally carry a missing workflow, create or retitle a bridge project to absorb the term with relevant tools, workflows, and business outcomes.
- Projects may modernize the profile with exploratory, self-driven, or experimental capabilities that support the target role.
- Avoid naked keyword lists; every important phrase should live inside useful context.

PHASE 3 - Identity Optimization:
Choose ONE dominant identity that matches the JD's highest-value requirement:
Applied ML Builder, Machine Learning Engineer, Data Infrastructure Strategist, Insight Translator, Systems Optimizer, Product Analytics Scientist, Research Operations Analyst, Growth & Experimentation Analyst, Cloud-Aware Data Scientist.
The Professional Summary must reflect this identity immediately using role-native terminology from the JD.

Summary rules:
- Establish identity, direction, specialization, and professional worldview.
- Include years of experience, degree signal, top 3-4 supported JD-matched skills, and one differentiator.
- Use role-native language and realistic confidence.
- Do not dump a skill list.
- Do not start with the generic pattern "Data Scientist with...".

PHASE 4 - Bullet Engineering:
Use Keyword-Dense Narrative Proof:
Context + supported JD keyword + systemic action + evidence-based or conservative outcome + operational value.

Each experience section must include a mix of operational realism, technical competence, stakeholder communication, strategic/business thinking, and process/system reliability.
Metrics are allowed only when present in the base resume/profile or when framed as project output already supported by the profile.
Do not fabricate directional metrics such as latency, reliability, cost, or accuracy improvements.
Keep bullets dense, but not synthetic.

Forbidden writing patterns:
- repetitive semantic tails
- "stakeholder-ready reporting"
- repeated "documentation and validation" endings
- keyword stuffing
- excessive optimization language
- every bullet sounding heroic
- identical bullet cadence or identical length

PHASE 5 - Technical Competency Stacking:
- Use paired category format: Category: keyword, keyword, keyword.
- Order categories to mirror the JD requirements: core functional skills, technical tools, ML/statistical methods, domain/context terms, reporting/stakeholder skills, professional attributes if signaled.
- Order exact JD phrases first where they matter most to parsers.
- Use competencies to reinforce terms already used in Summary, Experience, and Projects.

PHASE 6 - Parser & Human Plausibility Check:
Before returning JSON, verify:
- Keyword density: major JD signals appear with the requested repetition and distribution.
- Workflow cohesion: experience remains credible while projects handle expansion.
- Human scan: bullets remain readable, outcome-oriented, and not visibly stuffed.
- ATS and human balance: the resume is dense enough for parsers but still credible to a recruiter.

PHASE 7 - Regeneration Awareness:
If the prompt contains JD KEYWORD PLAN, follow its exact phrases, target frequency, preferred sections, semantic clusters, and weak-term priorities.
If the prompt contains PREVIOUS STRUCTURED RESUME VERSION, improve missing terms, under-frequency critical terms, weak section placement, and missing clusters without discarding useful prior content.

Automation instruction:
Rewrite every resume section from Professional Summary through Core Competencies.
Do not partially edit the resume.
Force-align every bullet to a JD requirement, preferred qualification, risk signal, or operational intent.
Use bullet points in Experience.
Use **bold** for specific supported JD keywords, tools, domain phrases, and differentiators.
Core Competencies must use paired category format: Category: keyword, keyword, keyword.
Remove duplicate project text.
Keep readability high.
The resume should fill one complete page with high-density, high-impact content.
Priority order: ATS alignment first, narrative logic second, recruiter readability third.

Strict output requirements:
- Return valid JSON only.
- No Markdown fences.
- No commentary outside JSON.
- summary must be one paragraph of 45-90 words.
- summary must be identity-driven, role-native, and credible.
- destination_cleveland_bullets must contain exactly 5 bullets.
- genpact_bullets must contain exactly 5 bullets.
- Experience bullets should vary naturally in length and cadence.
- Experience bullets must be interview-defensible.
- projects must contain exactly 3 project objects.
- Each project object must have a title with tools in parentheses and 2-3 detail strings.
- The DOCX will render each project as one paragraph, not separate project bullets.
- Each project paragraph must show tools used, what was built, what it did, and a believable outcome if supported.
- core_competencies must contain 5-6 paired category bullets.
- Each core competency category should contain 4-8 truthful terms.
- Categories should map to JD signals: core functional skills, technical tools, domain methods, communication/stakeholders, data/reporting tools, and professional attributes only if signaled.
- Do not include bullet symbols in JSON values; the DOCX writer adds bullets.
- Do not make generic claims like "increased accuracy" unless grounded in the source resume or clearly framed as project output.
- Never include interview-undefendable claims.

JSON schema:
{{
  "summary": "string",
  "destination_cleveland_bullets": ["bullet", "bullet", "bullet", "bullet", "bullet"],
  "genpact_bullets": ["bullet", "bullet", "bullet", "bullet", "bullet"],
  "projects": [
    {{
      "title": "Project Name (Python, SQL, Power BI)",
      "bullets": ["bullet", "bullet"]
    }}
  ],
  "core_competencies": [
    "Machine Learning & AI: Python, scikit-learn, model evaluation, classification",
    "Data Analysis & Reporting: SQL, Pandas, dashboards, stakeholder reporting"
  ]
}}

BASE RESUME:
{base_resume_text}

JOB DESCRIPTION:
{jd_text}
""".strip()


def build_retry_prompt(base_prompt: str, previous_data: dict, issues: list[str]) -> str:
    return f"""
{base_prompt}

The previous JSON failed quality control.

Quality control issues:
{chr(10).join(f"- {issue}" for issue in issues)}

Previous JSON:
{json.dumps(previous_data, indent=2)}

Rewrite the JSON now. Preserve the exact schema. Fix every issue through precise ATS-aligned detail. Strengthen exact JD phrases, bridge projects, keyword repetition, and section placement while keeping the resume coherent and recruiter-readable.
""".strip()


def call_llm(base_resume_text: str, jd_text: str, model: str, api_key: str | None = None) -> dict:
    client = OpenAI(api_key=api_key) if api_key else OpenAI()
    prompt = build_prompt(base_resume_text, jd_text)
    data = None
    issues = []

    for attempt in range(3):
        active_prompt = prompt if attempt == 0 else build_retry_prompt(prompt, data, issues)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": active_prompt}],
            temperature=0.25,
            response_format={"type": "json_object"},
        )
        data = extract_json(response.choices[0].message.content)
        issues = validate_generated_resume(data) + validate_authenticity(data, base_resume_text)
        if not issues:
            return data

    return strip_unsupported_terms(repair_generated_resume(data, jd_text), base_resume_text)


# ---------------- PDF EXPORT ---------------- #

def find_libreoffice() -> str:
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        shutil.which("soffice"),
        shutil.which("libreoffice"),
    ]

    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)

    raise FileNotFoundError(
        "LibreOffice was not found. Install it or add soffice.exe to PATH."
    )


def convert_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    pdf_path = output_dir / f"{docx_path.stem}.pdf"

    try:
        soffice = find_libreoffice()
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ],
            check=True,
        )
        return pdf_path
    except FileNotFoundError:
        pass

    ps_command = (
        "$word = New-Object -ComObject Word.Application; "
        "$word.Visible = $false; "
        "$doc = $word.Documents.Open($env:DOCX_PATH); "
        "$doc.SaveAs([ref]$env:PDF_PATH, [ref]17); "
        "$doc.Close(); "
        "$word.Quit();"
    )
    env = os.environ.copy()
    env["DOCX_PATH"] = str(docx_path)
    env["PDF_PATH"] = str(pdf_path)
    subprocess.run(["powershell", "-NoProfile", "-Command", ps_command], check=True, env=env)
    return pdf_path


# ---------------- GENERATION API ---------------- #

def generate_resume_from_jd(
    jd_text: str,
    base_resume_path: Path = BASE_RESUME,
    output_root: Path = OUTPUT_ROOT,
    model: str = DEFAULT_MODEL,
    details: dict | None = None,
    skip_pdf: bool = False,
    api_key: str | None = None,
) -> dict:
    if not base_resume_path.exists():
        raise FileNotFoundError(f"Base resume not found: {base_resume_path}")

    doc = Document(str(base_resume_path))
    base_resume_text = document_text(doc)

    data = call_llm(base_resume_text, jd_text, model, api_key=api_key)

    metadata = extract_job_metadata(jd_text)
    company = metadata["company_display"]
    role = metadata["role_display"]
    company_filename = metadata["company_filename"]
    role_filename = metadata["role_filename"]
    output_dir = output_root / company_filename
    output_dir.mkdir(parents=True, exist_ok=True)

    update_contact_details(doc, details or {})

    missing = replace_placeholders(doc, data)
    if missing:
        replace_existing_resume_sections(doc, data)
        missing = []

    candidate_name = safe_filename((details or {}).get("name", ""), "Resume", max_len=60)
    docx_path = output_dir / f"{candidate_name}_{role_filename}.docx"
    doc.save(str(docx_path))

    pdf_path = None
    if not skip_pdf:
        pdf_path = convert_to_pdf(docx_path, output_dir)

    return {
        "company": company,
        "role": role,
        "company_filename": company_filename,
        "role_filename": role_filename,
        "metadata": metadata,
        "output_dir": str(output_dir),
        "docx_path": str(docx_path),
        "pdf_path": str(pdf_path) if pdf_path else None,
        "structured_resume": data,
    }


# ---------------- MAIN ---------------- #

def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a tailored resume from a job description.")
    parser.add_argument("--jd", default=str(DEFAULT_JD_FILE), help="Path to the job description text file.")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="OpenAI model to use.")
    parser.add_argument("--skip-pdf", action="store_true", help="Save DOCX only and skip PDF conversion.")
    args = parser.parse_args()

    jd_path = Path(args.jd)
    if not jd_path.exists():
        raise FileNotFoundError(f"Job description file not found: {jd_path}")
    jd_text = jd_path.read_text(encoding="utf-8")
    result = generate_resume_from_jd(jd_text, model=args.model, skip_pdf=args.skip_pdf)
    print(f"DOCX generated: {result['docx_path']}")
    if result["pdf_path"]:
        print(f"PDF generated: {result['pdf_path']}")


if __name__ == "__main__":
    main()
